from selenium import webdriver
import docx
from datetime import datetime
from selenium.webdriver.common.keys import Keys
import random

i=0
links_NM=['https://wn.test1.nmg/','https://wn.test2.nmg/','https://wn.test3.nmg/','https://wn.test4.nmg/']               
links_BG=['https://wb.test1.nmg/','https://wb.test2.nmg/','https://wb.test3.nmg/','https://wb.test4.nmg/']
links_LC=['https://lc.test1.nmg/','https://lc.test2.nmg/','https://lc.test3.nmg/','https://lc.test4.nmg/']
links_HC=['https://wh.test1.nmg/','https://wh.test2.nmg/','https://wh.test3.nmg/','https://wh.test4.nmg/']

searchList_NM_BG_LC=['shoes','beauty','shirts','tshirts','boots','jewelry','women','men','bags','home','kids','sale']
searchList_HC=['bath','sale','bed','rugs','sheets','lightning','furniture','office','gifts']

now=datetime.now()
filename=now.strftime("%m%d%Y%H%M")

chrome_options = webdriver.ChromeOptions()
chrome_options.add_argument("--incognito")
chrome_options.add_argument('--ignore-certificate-errors')

doc=docx.Document()
filepath = 'C:\\Users\\Hemachandar\\Downloads\\tdr'+filename+'.docx' 
doc.save(filepath)

browser=webdriver.Chrome(executable_path='C:/Users/Hemachandar/Downloads/chromedriver.exe',chrome_options=chrome_options);
browser.maximize_window()

def searchScreenshot_NM(x):
    try:
        if(x=="https://wn.test4.nmg/"):
            search=browser.find_element_by_id("searchInput")
        else:
            search=browser.find_element_by_id("brSearchInput")  
                        
        search.send_keys(searchList_NM_BG_LC[random.randint(0,len(searchList_NM_BG_LC)-1)])
        search.send_keys(Keys.ENTER)
        #browser.find_element_by_xpath("//*[@id='search-form']/table/tbody/tr/td/input[1]").click()
    
    except Exception as e:
        print(e)
        pass
        
    links= browser.find_elements_by_class_name("product-thumbnail-image")
       
    return links;

def searchScreenshot_BG(x):
    try:
        if(x=="https://wb.test4.nmg/"):
            search=browser.find_element_by_id("searchInput")
            search.send_keys(searchList_NM_BG_LC[random.randint(0,len(searchList_NM_BG_LC)-1)])
            search.send_keys(Keys.ENTER)
        else:
            search=browser.find_elements_by_id("searchInput")       
            search[1].send_keys(searchList_NM_BG_LC[random.randint(0,len(searchList_NM_BG_LC)-1)])
            search[1].send_keys(Keys.ENTER)
    except Exception as e:
        print(e)
        pass
    
    links= browser.find_elements_by_class_name("product-thumbnail-image")   
    return links;

def searchScreenshot_LC(x):
    try:
        search=browser.find_element_by_id("searchInput")     
        search.send_keys(searchList_NM_BG_LC[random.randint(0,len(searchList_NM_BG_LC)-1)])
        search.send_keys(Keys.ENTER)
    
    except Exception as e:
        print(e)
        pass
        
    links= browser.find_elements_by_class_name("product-thumbnail-image")  
    return links;

def searchScreenshot_HC(x):      
    try:
        search=browser.find_element_by_id("searchInput")
        search.send_keys(searchList_HC[random.randint(0,len(searchList_HC)-1)])
        search.send_keys(Keys.ENTER)
                 
    except Exception as e:
        print(e)
        pass
    
    links=browser.find_elements_by_class_name("product-thumbnail-image")
    return links
   
    
for x in links_NM:
    i=i+1    
    browser.get(x) 
    browser.get(x)
    linksList=searchScreenshot_NM(x);  
    
    if len(linksList)>0:
        linksList[random.randint(0, len(linksList)-1)].click()  
                     
    name=browser.current_url[8:16]
    doc.add_paragraph(browser.current_url[8:16])
    doc.add_paragraph("Full URL: ")
    doc.add_paragraph(browser.current_url)             
    browser.save_screenshot("C:\\Users\\Hemachandar\\Downloads\\tdr\\" + str(i) + "_screenshot.png")             
    filename='C:\\Users\\Hemachandar\\Downloads\\tdr\\' + str(i)+ '_screenshot.png'             
    doc.add_picture(filename,width=docx.shared.Inches(7.5),height=docx.shared.Inches(4.5))
    doc.save(filepath)
        
for x in links_BG:
    i=i+1    
    browser.get(x)
    browser.get(x) 
    linksList=searchScreenshot_BG(x);  
    
    if len(linksList)>0:
        linksList[random.randint(0, len(linksList)-1)].click()  
                       
    name=browser.current_url[8:16]
    doc.add_paragraph(browser.current_url[8:16])
    doc.add_paragraph("Full URL: ")
    doc.add_paragraph(browser.current_url)             
    browser.save_screenshot("C:\\Users\\Hemachandar\\Downloads\\tdr\\" + str(i) + "_screenshot.png")             
    filename='C:\\Users\\Hemachandar\\Downloads\\tdr\\' + str(i)+ '_screenshot.png'             
    doc.add_picture(filename,width=docx.shared.Inches(7.5),height=docx.shared.Inches(4.5))
    doc.save(filepath)
        
for x in links_LC:
    i=i+1    
    browser.get(x)
    browser.get(x) 
    linksList=searchScreenshot_LC(x);  
    
    if len(linksList)>0:
        linksList[random.randint(0, len(linksList)-1)].click()  
                       
    name=browser.current_url[8:16]
    doc.add_paragraph(browser.current_url[8:16])
    doc.add_paragraph("Full URL: ")
    doc.add_paragraph(browser.current_url)             
    browser.save_screenshot("C:\\Users\\Hemachandar\\Downloads\\tdr\\" + str(i) + "_screenshot.png")             
    filename='C:\\Users\\Hemachandar\\Downloads\\tdr\\' + str(i)+ '_screenshot.png'             
    doc.add_picture(filename,width=docx.shared.Inches(7.5),height=docx.shared.Inches(4.5))
    doc.save(filepath)

for x in links_HC:
    i=i+1    
    browser.get(x)
    browser.get(x) 
    linksList=searchScreenshot_HC(x);  
    
    if len(linksList)>0:
        linksList[random.randint(0, len(linksList)-1)].click()  
                       
    name=browser.current_url[8:16]
    doc.add_paragraph(browser.current_url[8:16])
    doc.add_paragraph("Full URL: ")
    doc.add_paragraph(browser.current_url)             
    browser.save_screenshot("C:\\Users\\Hemachandar\\Downloads\\tdr\\" + str(i) + "_screenshot.png")             
    filename='C:\\Users\\Hemachandar\\Downloads\\tdr\\' + str(i)+ '_screenshot.png'             
    doc.add_picture(filename,width=docx.shared.Inches(7.5),height=docx.shared.Inches(4.5))
    doc.save(filepath)
    
browser.close()



